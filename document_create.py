import os
from docx import Document
from docx.shared import Inches
import folder_managnment
import string_manupulation
import collections

def AddpictureToWordFile(targetfile):
    document = Document()
    folderPath=folder_managnment.folder
    dictionaryPic={}
    document.add_paragraph()
    # r = p.add_run()
    with os.scandir(folderPath) as entries:
        for entry in entries:
            # print(entry.name)
            #r.add_text(text_above)
            print(entry.name)
            dictionaryPic[string_manupulation.FindNumber(entry.name)]=entry.name
    od=collections.OrderedDict(sorted(dictionaryPic.items()))
    for k, v in od.items():
        print(str(k) +" " + v)
        #document.add_picture(folderPath+ "\\" + item)
        #document.add_paragraph(entry.name.replace("Step","Figure"))

    document.save(targetfile)
    